"""
resume_parser.py

Reads a PDF or DOCX resume, extracts plain text (PyMuPDF for PDF,
python-docx for DOCX), and heuristically pulls out skills, technologies,
projects, and education. Prints a single JSON object to stdout so the
Node backend can spawn this as a subprocess and read the result.

Usage:
    python3 resume_parser.py <file_path> <file_type>
    file_type is one of: pdf, docx
"""

import sys
import json
import re

from resume_keywords import (
    TECHNOLOGIES,
    SKILLS,
    SECTION_HEADINGS,
    DEGREE_PATTERNS,
    EDUCATION_KEYWORDS,
)


def fail(message):
    print(json.dumps({"success": False, "error": message}))
    sys.exit(1)


def extract_text_from_pdf(path):
    import fitz  # PyMuPDF

    text_parts = []
    with fitz.open(path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def extract_text_from_docx(path):
    import docx

    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def find_matches(text, bank):
    """Case-insensitive whole-word-ish match of a keyword bank against text.
    Returns matches in the bank's original casing, deduped, longest-first
    so e.g. 'Node.js' is preferred over a bare 'Node' substring collision.
    """
    found = []
    lower_text = text.lower()
    for term in sorted(bank, key=len, reverse=True):
        pattern = r"(?<![a-zA-Z0-9])" + re.escape(term.lower()) + r"(?![a-zA-Z0-9])"
        if re.search(pattern, lower_text):
            found.append(term)
    # Dedupe while preserving first-seen order, case-insensitively.
    seen = set()
    deduped = []
    for term in found:
        key = term.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(term)
    return deduped


def split_into_sections(text):
    """Split resume text into {section_name: body_text} using known headings.
    A line counts as a heading if, once stripped and uppercased, it matches
    (or closely matches) one of the known heading phrases and is short.
    """
    lines = [line.strip() for line in text.splitlines()]
    heading_lookup = {}
    for section, phrases in SECTION_HEADINGS.items():
        for phrase in phrases:
            heading_lookup[phrase.upper()] = section

    sections = {}
    current_section = "header"
    buffer = []

    for line in lines:
        clean = re.sub(r"[^A-Za-z& ]", "", line).strip().upper()
        matched_section = None
        if 0 < len(clean) <= 40 and clean in heading_lookup:
            matched_section = heading_lookup[clean]

        if matched_section:
            sections[current_section] = "\n".join(buffer).strip()
            current_section = matched_section
            buffer = []
        else:
            buffer.append(line)

    sections[current_section] = "\n".join(buffer).strip()
    return sections


def parse_skills_section(section_text):
    if not section_text:
        return []
    # Skill sections are usually comma/bullet/newline separated tokens.
    tokens = re.split(r"[\n,;|•·]+", section_text)
    cleaned = []
    for token in tokens:
        token = token.strip(" -\t")
        # Drop category labels like "Languages:" leaving just the value side.
        if ":" in token:
            token = token.split(":", 1)[1].strip()
        if 1 < len(token) <= 40:
            cleaned.append(token)
    return cleaned


def parse_projects_section(section_text, full_text_technologies_bank):
    if not section_text:
        return []

    # Split on blank lines first; if that yields one giant block, fall back
    # to splitting on lines that look like bullet/title starts.
    blocks = [b.strip() for b in re.split(r"\n\s*\n", section_text) if b.strip()]
    if len(blocks) <= 1:
        blocks = [b.strip() for b in re.split(r"\n(?=[A-Z\u2022\-\*])", section_text) if b.strip()]

    projects = []
    for block in blocks[:12]:  # sane upper bound
        lines = [l.strip(" -•\t") for l in block.splitlines() if l.strip()]
        if not lines:
            continue
        title = lines[0][:120]
        description_lines = lines[1:]
        description = " ".join(description_lines)[:600]
        technologies = find_matches(block, full_text_technologies_bank)
        projects.append(
            {
                "title": title,
                "description": description,
                "technologies": technologies,
            }
        )
    return projects


def parse_education_section(section_text):
    if not section_text:
        return []

    blocks = [b.strip() for b in re.split(r"\n\s*\n", section_text) if b.strip()]
    if len(blocks) <= 1:
        # Fall back to one block per line-group around institution keywords.
        blocks = [section_text]

    education = []
    for block in blocks[:6]:
        lines = [l.strip(" -•\t") for l in block.splitlines() if l.strip()]
        if not lines:
            continue

        degree = None
        institution = None
        for line in lines:
            if degree is None:
                for pattern in DEGREE_PATTERNS:
                    if pattern.lower() in line.lower():
                        degree = line[:120]
                        break
            if institution is None:
                for kw in EDUCATION_KEYWORDS:
                    if kw.lower() in line.lower():
                        institution = line[:120]
                        break

        year_match = re.search(r"(19|20)\d{2}", block)
        year = year_match.group(0) if year_match else None

        if degree or institution:
            education.append(
                {
                    "degree": degree or "Degree not detected",
                    "institution": institution or "Institution not detected",
                    "year": year,
                }
            )

    return education


def parse_resume(path, file_type):
    if file_type == "pdf":
        text = extract_text_from_pdf(path)
    elif file_type == "docx":
        text = extract_text_from_docx(path)
    else:
        fail(f"Unsupported file type: {file_type}")
        return

    if not text or not text.strip():
        fail("Could not extract any text from this file. It may be a scanned/image-only document.")
        return

    sections = split_into_sections(text)

    skills_section_text = sections.get("skills", "")
    explicit_skills = parse_skills_section(skills_section_text)

    # Technologies: curated bank matched across the *entire* resume, so tools
    # mentioned in project descriptions are captured even outside a Skills section.
    technologies = find_matches(text, TECHNOLOGIES)

    # Broader competency/practice keywords (System Design, TDD, etc.)
    soft_skills = find_matches(text, SKILLS)

    # Merge explicit skill-section tokens with curated technology matches,
    # keeping the result readable rather than dumping every raw token.
    skills = explicit_skills if explicit_skills else technologies

    projects = parse_projects_section(sections.get("projects", ""), TECHNOLOGIES)
    education = parse_education_section(sections.get("education", ""))

    result = {
        "success": True,
        "data": {
            "skills": skills[:30],
            "technologies": technologies[:30],
            "softSkills": soft_skills[:15],
            "projects": projects,
            "education": education,
            "textLength": len(text),
        },
    }
    print(json.dumps(result))


if __name__ == "__main__":
    if len(sys.argv) < 3:
        fail("Usage: python3 resume_parser.py <file_path> <file_type>")
    file_path_arg = sys.argv[1]
    file_type_arg = sys.argv[2].lower().lstrip(".")
    try:
        parse_resume(file_path_arg, file_type_arg)
    except Exception as exc:  # noqa: BLE001 - report any failure back as JSON
        fail(f"Failed to parse resume: {exc}")
